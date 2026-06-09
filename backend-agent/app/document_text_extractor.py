from __future__ import annotations

import csv
import io
import json
import re
from io import BytesIO

from app.exceptions import ValidationError


def markdown_to_plain_text(md_content: str) -> str:
    md_content = re.sub(r"```[\s\S]*?```", "", md_content)
    md_content = re.sub(r"`[^`]+`", "", md_content)
    md_content = re.sub(r"!\[.*?\]\(.*?\)", "", md_content)
    md_content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", md_content)
    md_content = re.sub(r"\*\*([^*]+)\*\*", r"\1", md_content)
    md_content = re.sub(r"__([^_]+)__", r"\1", md_content)
    md_content = re.sub(r"\*([^*]+)\*", r"\1", md_content)
    md_content = re.sub(r"_([^_]+)_", r"\1", md_content)
    md_content = re.sub(r"^#{1,6}\s*", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^>\s*", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^[-*+]\s*", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\d+\.\s*", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^[-*_]{3,}\s*$", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"<[^>]+>", "", md_content)
    md_content = re.sub(r"\n\s*\n", "\n\n", md_content).strip()
    return md_content


def decode_bytes(content: bytes, label: str = "文件") -> str:
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValidationError(f"无法识别{label}编码")


def extract_text_from_txt(content: bytes) -> str:
    return decode_bytes(content, "文本文件")


def extract_text_from_md(content: bytes) -> str:
    raw = decode_bytes(content, "Markdown 文件")
    return markdown_to_plain_text(raw)


def extract_text_from_json(content: bytes) -> str:
    raw = decode_bytes(content, "JSON 文件")
    try:
        parsed = json.loads(raw)
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        pass
    for wrapper in [("[", "]"), ("{\n", "}\n")]:
        if not raw.strip().startswith(wrapper[0][0]):
            try:
                parsed = json.loads(wrapper[0] + raw + wrapper[1])
                return json.dumps(parsed, ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                continue
    return raw


def extract_text_from_csv(content: bytes) -> str:
    raw = decode_bytes(content, "CSV 文件")
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return ""
    lines = []
    for row in rows:
        lines.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(lines)


def _extract_table_text(table) -> str:
    rows_text = []
    for row in table.rows:
        cells_text = []
        for cell in row.cells:
            cell_text = "\n".join(
                para.text for para in cell.paragraphs if para.text.strip()
            )
            cells_text.append(cell_text.strip())
        rows_text.append(" | ".join(cells_text))
    return "\n".join(rows_text)


def extract_text_from_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(BytesIO(content))
        parts = []
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        table_text = _extract_table_text(table)
                        if table_text:
                            parts.append(table_text)
                        break
        return "\n".join(parts)
    except ImportError:
        return ""


def extract_text_from_doc(content: bytes) -> str:
    try:
        import olefile
    except ImportError:
        return ""
    try:
        ole = olefile.OleFileIO(BytesIO(content))
        if not ole.exists('WordDocument'):
            return ""
        word_stream = ole.openstream('WordDocument').read()
        wIdent = int.from_bytes(word_stream[0:2], 'little')
        if wIdent != 0xA5EC:
            return ""
        flags = int.from_bytes(word_stream[0x000A:0x000C], 'little')
        table_name = '1Table' if (flags & 0x0200) else '0Table'
        if not ole.exists(table_name):
            return ""
        table_stream = ole.openstream(table_name).read()
        pos = 32
        csw = int.from_bytes(word_stream[pos:pos + 2], 'little')
        pos += 2 + csw * 2
        cslw = int.from_bytes(word_stream[pos:pos + 2], 'little')
        pos += 2 + cslw * 4
        pos += 2
        fcClx_idx = pos + 33 * 8
        if fcClx_idx + 8 > len(word_stream):
            return ""
        fcClx = int.from_bytes(word_stream[fcClx_idx:fcClx_idx + 4], 'little')
        lcbClx = int.from_bytes(word_stream[fcClx_idx + 4:fcClx_idx + 8], 'little')
        if lcbClx == 0 or fcClx + lcbClx > len(table_stream):
            return ""
        clx_data = table_stream[fcClx:fcClx + lcbClx]
        clx_pos = 0
        piece_table = b''
        while clx_pos < len(clx_data):
            marker = clx_data[clx_pos]
            if marker == 0x02:
                clx_pos += 1
                pt_size = int.from_bytes(clx_data[clx_pos:clx_pos + 4], 'little')
                clx_pos += 4
                piece_table = clx_data[clx_pos:clx_pos + pt_size]
                break
            elif marker == 0x01:
                clx_pos += 1
                cb = int.from_bytes(clx_data[clx_pos:clx_pos + 2], 'little')
                clx_pos += 2 + cb
            else:
                break
        if not piece_table:
            return ""
        n = (len(piece_table) - 4) // 12
        if n <= 0:
            return ""
        cps = [int.from_bytes(piece_table[i * 4:(i + 1) * 4], 'little') for i in range(n + 1)]
        pcd_start = (n + 1) * 4
        text_parts = []
        for i in range(n):
            pcd = piece_table[pcd_start + i * 8:pcd_start + (i + 1) * 8]
            fc_raw = int.from_bytes(pcd[2:6], 'little')
            is_compressed = bool(fc_raw & 0x40000000)
            fc = fc_raw & ~0x40000000
            char_count = cps[i + 1] - cps[i]
            if char_count <= 0:
                continue
            if is_compressed:
                start = fc // 2
                raw = word_stream[start:start + char_count]
                text_parts.append(raw.decode('cp1252', errors='replace'))
            else:
                raw = word_stream[fc:fc + char_count * 2]
                text_parts.append(raw.decode('utf-16-le', errors='replace'))
        full_text = ''.join(text_parts)
        full_text = full_text.replace('\x00', '').replace('\r\n', '\n').replace('\r', '\n')
        lines = [line.strip() for line in full_text.split('\n')]
        return '\n'.join(line for line in lines if line)
    except Exception:
        return ""


EXTRACTORS = {
    ".txt": extract_text_from_txt,
    ".md": extract_text_from_md,
    ".json": extract_text_from_json,
    ".csv": extract_text_from_csv,
    ".docx": extract_text_from_docx,
    ".doc": extract_text_from_doc,
}


def extract_text_from_file(content: bytes, ext: str) -> str:
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        return ""
    try:
        return extractor(content)
    except Exception:
        return ""


def validate_characters(text: str, filename: str, max_characters: int) -> None:
    if text:
        char_count = len(text)
        if char_count > max_characters:
            raise ValidationError(
                f"文件 {filename} 字符数 {char_count} 超过限制（最大 {max_characters}）。"
                f"建议将文档按章节或主题手动切分为多个小文件，命名格式为后缀增加_1 _2，例如 说明书_1.doc 说明书_2.doc，然后一次性上传。"
                f"手动切分有助于提升记忆的准确性和完整性。"
            )
